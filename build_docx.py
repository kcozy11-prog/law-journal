"""
Hearts For My Sweet Baboo — 영어 학습지를 인쇄용 Word(.docx) 파일로 생성.

웹 앱(snoopy-valentine.html)과 동일한 스크립트·빈칸·문법 패턴을 그대로 사용.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 데이터 (snoopy-valentine.html 과 동일) ─────────────
SCENES = [
    {
        "id": 1,
        "title": "Scene 1 — Charlie Brown and Snoopy with Valentines",
        "title_ko": "장면 1 — 찰리 브라운과 스누피, 그리고 발렌타인 카드",
        "lines": [
            ("Charlie Brown",  "Snoopy, look at all these {Valentines}!", "스누피, 이 발렌타인 카드들 좀 봐!", None),
            ("Charlie Brown",  "There's even one for you.", "네 것도 있어.", None),
            ("Charlie Brown",  "How about that?", "어때?", None),
            ("Charlie Brown",  "Oh, here's {another} one for you.", "오, 여기 또 하나 더 있네.", None),
            ("Charlie Brown",  "And another!", "또 있어!", None),
            ("Charlie Brown",  "{Must be} your lucky day.", "오늘은 운수 좋은 날인가 본데.", "MODAL_GUESS"),
            ("Charlie Brown",  "One {more} for you.", "여기 하나 더.", None),
            ("Charlie Brown",  "One more for you.", "하나 더.", None),
            ("Charlie Brown",  "Wait, here's one with my {name} on it.", "잠깐, 여기 내 이름이 적힌 게 하나 있네.", None),
            ("Charlie Brown (reading)", '"Charlie Brown, please {give this to} Snoopy."', '"찰리 브라운, 이것 좀 스누피에게 전해줘."', "GIVE_OBJ_TO"),
            ("Charlie Brown",  "I didn't even {get} one.", "난 한 장도 못 받았는데.", None),
            ("Charlie Brown (reading)", '"To the {round-headed kid who} brings me supper."', '"나에게 저녁을 가져다주는 머리가 둥근 꼬마에게."', "RELATIVE_WHO"),
            ("Charlie Brown (reading)", '"Happy Valentine\'s Day."', '"행복한 발렌타인 데이 보내."', None),
            ("Charlie Brown",  "I know I {should be} embarrassed, but I'll take it.", "창피해야 한다는 건 알지만, 그래도 받을게.", "MODAL_DUTY"),
            ("Charlie Brown",  "Thanks, old {pal}.", "고맙다 친구야.", None),
            ("Narrator",       "Charlie Brown sure {is excited about} getting mail.", "찰리 브라운은 우편물 받는 걸 정말 좋아하는구나.", "BE_EXCITED_ABOUT"),
        ],
    },
    {
        "id": 2,
        "title": "Scene 2 — Sally making a card",
        "title_ko": "장면 2 — 샐리의 카드 만들기",
        "lines": [
            ("Sally", "Hello, my sweet {baboo}!", "안녕, 나의 귀염둥이!", None),
            ("Sally", "{Nothing says} Valentine's Day {like} a {homemade} card.", "직접 만든 카드만큼 발렌타인 데이에 어울리는 건 없지.", "NOTHING_LIKE"),
            ("Sally", "Let's see...", "보자...", None),
            ("Sally", "I {need to add} some glue, sparkles, and the ultimate {finishing} touch:", "풀이랑 반짝이도 좀 넣고, 마지막 결정적인 한 방으로,", "NEED_TO_V"),
            ("Sally", "a macaroni {self-portrait}.", "마카로니 자화상을 붙여야지.", None),
            ("Sally", "Perfection!", "완벽해!", None),
            ("Sally", "Snoopy, I {need you to} take this.", "스누피, 이것 좀 가져가 줘.", "NEED_O_TO_V"),
            ("Sally", "It's not for you.", "네 거 아니야.", None),
            ("Sally", "I {need you to} deliver it.", "배달해 줘야 해.", "NEED_O_TO_V"),
            ("Sally", "After all, there's {nothing classier than} a personal courier.", "결국 개인 심부름꾼만큼 품격 있는 건 없으니까.", "NOTHING_THAN"),
            ("Sally", "Take it to my sweet baboo.", "내 귀염둥이에게 가져다줘.", None),
            ("Sally", "He's {the one who} puts a song in my heart.", "그는 내 마음에 노래를 불러주는 사람이야.", "RELATIVE_WHO"),
            ("Sally", "Got it?", "알겠지?", None),
        ],
    },
    {
        "id": 3,
        "title": "Scene 3 — Delivery mistakes",
        "title_ko": "장면 3 — 계속되는 배달 실수",
        "lines": [
            ("Schroeder", "Can I {help} you?", "도와줄까?", None),
            ("Schroeder", "This is for {me}?", "이게 내 거라고?", None),
            ("Sally",     "That {must be} my sweet baboo!", "내 귀염둥이임에 틀림없어!", "MODAL_GUESS"),
            ("Sally",     "I'll {bet} he really liked my macaroni portrait.", "내 마카로니 초상화를 정말 좋아했을 거야.", None),
            ("Sally",     "Hello?", "여보세요?", None),
            ("Sally",     "He did what?", "그가 뭘 했다고?", None),
            ("Sally",     "This is {outrageous}!", "이건 말도 안 돼!", None),
            ("Sally",     "I mean, I'm {glad you liked} it, Schroeder.", "내 말은, 네가 좋아했다니 다행이네, 슈로더.", "BE_ADJ_THAT"),
            ("Sally",     "Snoopy, you {gave my Valentine to} the wrong person!", "스누피, 내 발렌타인 카드를 엉뚱한 사람한테 줬잖아!", "GIVE_OBJ_TO"),
            ("Sally",     "Now I have to make another masterpiece {that captures} the true depth of my love.", "이제 내 사랑의 깊이를 담은 또 다른 걸작을 만들어야 하잖아.", "RELATIVE_THAT"),
            ("Sally",     "You think that's {easy}?", "이게 쉬운 줄 아니?", None),
            ("Sally",     "It's not.", "아니라고.", None),
            ("Sally",     "{Construction} paper, glue, {macaroni}, extra {sparkles}...", "도화지, 풀, 마카로니, 반짝이 추가...", None),
            ("Sally",     "Perfect!", "완벽해!", None),
        ],
    },
    {
        "id": 4,
        "title": "Scene 4 — Final resolution",
        "title_ko": "장면 4 — 마지막 결말",
        "lines": [
            ("Charlie Brown", "Sally, I'm {the last person to} give advice on Valentine's Day,", "샐리, 내가 발렌타인 데이에 대해 조언할 처지는 아니지만,", "LAST_TO_V"),
            ("Charlie Brown", "but {have you considered} that it {might not be} Snoopy's fault?", "이게 스누피 잘못이 아닐 수도 있다는 생각은 안 해봤니?", "PRESENT_PERFECT_Q"),
            ("Charlie Brown", "Maybe he wasn't {sure who} your sweet baboo is.", "아마 네 귀염둥이가 누구인지 잘 몰랐을 수도 있잖아.", "BE_ADJ_WH"),
            ("Sally",         "Well, he {could have just asked} me!", "그럼 그냥 나한테 물어봤으면 됐잖아!", "COULD_HAVE_PP"),
            ("Sally",         "It {doesn't matter} now.", "이제 상관없어.", None),
            ("Sally",         "Valentine's Day {is ruined}.", "발렌타인 데이는 완전히 망쳤어.", "PASSIVE"),
            ("Charlie Brown", "It's not your {fault}, Snoopy.", "네 잘못이 아니야 스누피.", None),
            ("Charlie Brown", "It's just Valentine's Day {means a lot to} Sally.", "그냥 샐리한테 발렌타인 데이가 의미가 커서 그래.", "MEANS_A_LOT"),
            ("Friends",       "Because you made {so many} nice Valentine's Day cards for everyone,", "네가 모두를 위해 예쁜 발렌타인 카드를 많이 만들어줘서,", None),
            ("Friends",       "we {wanted to do} something special for you.", "우리도 널 위해 특별한 걸 준비하고 싶었어.", "WANT_TO_V"),
            ("Friends",       "Happy Valentine's Day!", "해피 발렌타인 데이!", None),
            ("Sally",         "For me?", "내 거니?", None),
            ("Sally",         "Oh, thank you!", "오, 고마워!", None),
            ("Sally",         "You even got the macaroni-to-sparkle {ratio} just right.", "마카로니와 반짝이 비율을 아주 딱 맞게 맞췄구나.", None),
            ("Charlie Brown", "Huh, {who knew} a day in the middle of February", "허, 2월 한복판의 어느 하루가", "WHO_KNEW"),
            ("Charlie Brown", "could bring so many people {together}?", "이렇게 많은 사람을 하나로 모을 줄 누가 알았겠어?", None),
        ],
    },
]


GRAMMAR = {
    "MODAL_GUESS":      ("조동사 추측",           "must / might + be",                 "~임에 틀림없다 (강한 확신) / ~일지도 모른다 (약한 추측)."),
    "MODAL_DUTY":       ("조동사 의무",           "should + be / V",                   "~해야 한다, ~하는 게 마땅하다."),
    "COULD_HAVE_PP":    ("가정법 과거완료",       "could have + p.p.",                 "~할 수도 있었는데 (실제로는 안 했다) — 후회·아쉬움."),
    "RELATIVE_WHO":     ("관계대명사 who",        "사람 + who + V",                    "앞의 사람을 꾸며줄 때 who 로 이어 붙인다."),
    "RELATIVE_THAT":    ("관계대명사 that",       "사물 + that + V",                   "사물·동물을 꾸며줄 때 that 으로 이어 붙인다."),
    "NEED_O_TO_V":      ("~에게 …하도록 하다",    "need + 사람 + to V",                 '"~가 ~해주기를 원/필요로 한다." want / ask / tell 도 같은 구조.'),
    "NEED_TO_V":        ("~할 필요가 있다",       "need to + V",                       "주어가 직접 ~할 필요가 있다."),
    "NOTHING_LIKE":     ("Nothing ~ like",       "Nothing says A like B",             '"B만큼 A를 잘 표현하는 건 없다" — 강조 구문.'),
    "NOTHING_THAN":     ("Nothing 비교급 than",   "Nothing 비교급 + than",             '"~보다 더 ~한 것은 없다" → 사실상 최상급 의미.'),
    "PRESENT_PERFECT_Q":("현재완료 의문문",       "Have you (ever) + p.p. ?",          '"~해본 적 있니? / ~해봤니?" — 경험·완료를 묻는다.'),
    "PASSIVE":          ("수동태",               "be + p.p.",                          "주어가 동작을 당하는 표현."),
    "LAST_TO_V":        ("the last + 명사 + to V","the last person to V",              '"~할 사람이 결코 아니다" — 반어법.'),
    "BE_ADJ_THAT":      ("be + 감정 형용사 + (that) 절", "I'm glad / sorry (that) S V", "감정 형용사 뒤에 절이 옴. that 은 자주 생략."),
    "BE_ADJ_WH":        ("be + 형용사 + 의문사절","be sure who / what / where ...",    '"누가/무엇이 …인지 확신한다(아니다)".'),
    "WHO_KNEW":         ("수사 의문문 — Who knew", "Who knew (that) … ?",              '"누가 ~할 줄 알았겠어!" — 놀라움·감탄.'),
    "WANT_TO_V":        ("want + to V",          "want / wanted + to V",              "~하고 싶다 / ~하고 싶었다."),
    "GIVE_OBJ_TO":      ("give + 사물 + to + 사람", "give A to B",                     '"B에게 A를 주다." 4형식 → 3형식.'),
    "BE_EXCITED_ABOUT": ("be + 감정 형용사 + about", "be excited / happy + about",     "~에 대해 신이 나다 / 들뜨다."),
    "MEANS_A_LOT":      ("mean a lot to ~",      "mean a lot to + 사람",              "~에게 큰 의미가 있다."),
}


# ── 유틸 ─────────────────────────────────────────────────
import re

CHUNK_RE = re.compile(r"\{([^}]+)\}")


def parse_chunks(text):
    """'I {need to} go' → [('text','I '), ('blank','need to'), ('text',' go')]"""
    parts = []
    last = 0
    for m in CHUNK_RE.finditer(text):
        if m.start() > last:
            parts.append(("text", text[last:m.start()]))
        parts.append(("blank", m.group(1)))
        last = m.end()
    if last < len(text):
        parts.append(("text", text[last:]))
    return parts


def all_blanks(scene):
    out = []
    for _who, en, _ko, _grammar in scene["lines"]:
        for kind, val in parse_chunks(en):
            if kind == "blank":
                out.append(val)
    return out


def filled_text(text):
    """빈칸 자리에 정답 그대로 넣어 완전한 영어 문장 반환."""
    return CHUNK_RE.sub(r"\1", text)


def make_blank_underline(chunk):
    """빈칸 길이에 맞춰 밑줄 ___________ 생성."""
    # 청크 글자수의 1.4배 정도, 최소 8칸
    return "_" * max(int(len(chunk) * 1.4) + 2, 8)


# ── 스타일 헬퍼 ──────────────────────────────────────────
def set_font(run, name=None, size=None, bold=None, italic=None,
             color=None, eastasia=None):
    if name:
        run.font.name = name
        if eastasia is None:
            eastasia = name
    if eastasia:
        # 한글에 동일 폰트 적용
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), eastasia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_hr(paragraph):
    """문단 아래에 가로선."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C4365C")
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


# ── 문서 작성 ────────────────────────────────────────────
PINK_DARK = (0xC4, 0x36, 0x5C)
PINK = (0xFF, 0x5C, 0x85)
INK = (0x2B, 0x1D, 0x24)
INK_SOFT = (0x5A, 0x46, 0x51)
GOLD = (0xD9, 0x9A, 0x4A)
BLUE = (0x1A, 0x5B, 0x94)
GREEN = (0x6A, 0x8F, 0x4D)


def make_doc(out_path):
    doc = Document()

    # 여백
    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rFonts.set(qn("w:ascii"), "Malgun Gothic")
    rFonts.set(qn("w:hAnsi"), "Malgun Gothic")

    # ── 표지 ────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PEANUTS · 스누피와 함께하는 영어")
    set_font(r, size=11, color=PINK_DARK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hearts For My Sweet Baboo ♥")
    set_font(r, size=26, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("발렌타인 데이 에피소드 영어 학습지")
    set_font(r, size=12, color=INK_SOFT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("이름: ____________________     날짜: ______________")
    set_font(r, size=11, color=INK_SOFT)

    p = doc.add_paragraph()
    add_hr(p)

    # 사용법 안내
    p = doc.add_paragraph()
    r = p.add_run("📖  사용 방법")
    set_font(r, size=13, bold=True, color=PINK_DARK)
    bullets = [
        "각 장면 위에 있는 단어 상자에서 알맞은 단어/구문을 골라 빈칸(________)에 적어 보세요.",
        "화자 이름 옆에 있는 ⚙ 표시는 문법 패턴 태그예요. 어떤 문법이 쓰였는지 미리 힌트를 줘요.",
        "마지막 페이지에는 같은 문법끼리 모은 \"패턴 연습\"과 정답키가 있어요.",
    ]
    for b in bullets:
        bp = doc.add_paragraph(style=None)
        bp.paragraph_format.left_indent = Cm(0.5)
        r = bp.add_run("• " + b)
        set_font(r, size=11)

    doc.add_page_break()

    # ── 각 장면 ──────────────────────────────────────────
    for scene in SCENES:
        # 장면 제목
        p = doc.add_paragraph()
        r = p.add_run(scene["title"])
        set_font(r, size=18, bold=True, color=PINK_DARK)

        p = doc.add_paragraph()
        r = p.add_run(scene["title_ko"])
        set_font(r, size=11, italic=True, color=INK_SOFT)

        add_hr(p)

        # 단어 상자
        chunks = list(dict.fromkeys(all_blanks(scene)))  # 순서 유지 + 중복 제거
        if chunks:
            box = doc.add_table(rows=1, cols=1)
            box.autofit = True
            cell = box.rows[0].cells[0]
            shade_cell(cell, "FFF8E7")

            # 헤더
            ph = cell.paragraphs[0]
            r = ph.add_run("단어 상자 · WORD BOX")
            set_font(r, size=10, bold=True, color=GOLD)

            # 칩들
            chip_p = cell.add_paragraph()
            for i, c in enumerate(chunks):
                if i:
                    sep = chip_p.add_run("   ·   ")
                    set_font(sep, color=GOLD, size=11, bold=True)
                rc = chip_p.add_run(c)
                set_font(rc, size=12, bold=True, color=INK)
            doc.add_paragraph()  # spacing

        # 대사 라인들
        for (who, en, ko, grammar_id) in scene["lines"]:
            # 화자 + 문법 태그
            wp = doc.add_paragraph()
            wp.paragraph_format.space_before = Pt(8)
            wp.paragraph_format.space_after = Pt(2)
            rwho = wp.add_run(who.upper())
            set_font(rwho, size=9, bold=True, color=PINK_DARK)
            if grammar_id and grammar_id in GRAMMAR:
                label, formula, _note = GRAMMAR[grammar_id]
                rsep = wp.add_run("   ⚙ ")
                set_font(rsep, size=9, color=BLUE)
                rg = wp.add_run(f"{label}  ·  {formula}")
                set_font(rg, size=9, bold=True, color=BLUE)

            # 영어 (빈칸 처리)
            ep = doc.add_paragraph()
            ep.paragraph_format.space_after = Pt(2)
            ep.paragraph_format.left_indent = Cm(0.4)
            for kind, val in parse_chunks(en):
                if kind == "text":
                    rt = ep.add_run(val)
                    set_font(rt, size=13, color=INK)
                else:
                    rb = ep.add_run(make_blank_underline(val))
                    set_font(rb, size=13, color=PINK_DARK, bold=True)

            # 한국어
            kp = doc.add_paragraph()
            kp.paragraph_format.space_after = Pt(4)
            kp.paragraph_format.left_indent = Cm(0.4)
            rk = kp.add_run("→ " + ko)
            set_font(rk, size=10, italic=True, color=INK_SOFT)

        doc.add_page_break()

    # ── 패턴 연습 ────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("패턴 연습  ·  Grammar Patterns")
    set_font(r, size=20, bold=True, color=BLUE)

    p = doc.add_paragraph()
    r = p.add_run("같은 문법 패턴을 가진 문장끼리 모았어요. 같은 구조가 반복되면 더 잘 외워져요.")
    set_font(r, size=11, color=INK_SOFT, italic=True)

    add_hr(p)

    # 그룹화
    grouped = {}
    for scene in SCENES:
        for (who, en, ko, gid) in scene["lines"]:
            if gid:
                grouped.setdefault(gid, []).append({
                    "scene_id": scene["id"], "who": who, "en": en, "ko": ko
                })

    for gid, (label, formula, note) in GRAMMAR.items():
        if gid not in grouped:
            continue
        examples = grouped[gid]

        # 패턴 헤더
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(12)
        rh = ph.add_run(label)
        set_font(rh, size=15, bold=True, color=BLUE)

        # 공식
        fp = doc.add_paragraph()
        fp.paragraph_format.space_after = Pt(2)
        rf = fp.add_run("[ " + formula + " ]")
        set_font(rf, size=11, bold=True, color=BLUE)

        # 설명
        np_ = doc.add_paragraph()
        np_.paragraph_format.space_after = Pt(6)
        rn = np_.add_run(note)
        set_font(rn, size=10, italic=True, color=INK_SOFT)

        # 예문
        for ex in examples:
            ep = doc.add_paragraph()
            ep.paragraph_format.left_indent = Cm(0.6)
            ep.paragraph_format.space_after = Pt(0)
            rl = ep.add_run(f"Scene {ex['scene_id']} · {ex['who']}\n")
            set_font(rl, size=8, color=PINK_DARK, bold=True)

            # 영어 (청크 강조)
            for kind, val in parse_chunks(ex["en"]):
                if kind == "text":
                    rt = ep.add_run(val)
                    set_font(rt, size=12, color=INK)
                else:
                    rb = ep.add_run(val)
                    set_font(rb, size=12, bold=True, color=GOLD)
                    # 밑줄
                    rb.underline = True

            kp = doc.add_paragraph()
            kp.paragraph_format.left_indent = Cm(0.6)
            kp.paragraph_format.space_after = Pt(8)
            rk = kp.add_run("→ " + ex["ko"])
            set_font(rk, size=9, italic=True, color=INK_SOFT)

    doc.add_page_break()

    # ── 정답키 ───────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("정답키  ·  Answer Key")
    set_font(r, size=20, bold=True, color=GREEN)

    p = doc.add_paragraph()
    r = p.add_run("스스로 풀어본 뒤에만 확인하세요!")
    set_font(r, size=11, italic=True, color=INK_SOFT)

    add_hr(p)

    for scene in SCENES:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(10)
        rs = sp.add_run(scene["title"])
        set_font(rs, size=12, bold=True, color=PINK_DARK)

        for i, (who, en, ko, gid) in enumerate(scene["lines"], 1):
            blanks = [v for k, v in parse_chunks(en) if k == "blank"]
            if not blanks:
                continue
            ap = doc.add_paragraph()
            ap.paragraph_format.left_indent = Cm(0.6)
            ap.paragraph_format.space_after = Pt(1)
            r1 = ap.add_run(f"{i}. ")
            set_font(r1, size=10, bold=True, color=INK_SOFT)

            # 완성 문장 (정답이 노란 배경 강조)
            for kind, val in parse_chunks(en):
                if kind == "text":
                    rt = ap.add_run(val)
                    set_font(rt, size=10, color=INK)
                else:
                    rt = ap.add_run(val)
                    set_font(rt, size=10, bold=True, color=GREEN)
                    rt.underline = True

    # 저장
    doc.save(out_path)
    print(f"✓ Saved: {out_path}")


if __name__ == "__main__":
    make_doc("/home/user/law-journal/snoopy-valentine.docx")
